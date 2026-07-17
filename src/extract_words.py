"""Farkli dosya turlerinden (kelime, tur) ciftleri cikarir.

Ogretmenlerin ne yukleyecegi belirsiz - Oxford 5000 gibi temiz tek-sutunlu
listeler de gelir, Cambridge PET gibi cok-sutunlu tematik gridler de. Ortak
kurtaric sinyal: bu listelerin hepsinde kelimenin yaninda bir TUR ISARETI var
(v., (n), (adj & prep) ...). Isaretten oncesi kelime, sonrasi (telaffuz, tanim,
ornek cumle) atilir.

Desteklenen: .txt, .pdf, .docx
"""
import os
import re

from src.read_words_from_txt import read_words_from_txt

# ── Tur isareti kaliplari ─────────────────────────────────────────────────
# Parantezli:  (n)  (n.)  (adj & prep)  (v., n.)  (phr.)
# Duz (satir sonu):  "absorb v."  "dive v., n."
_POS = r"adj|adv|prep|conj|pron|interj|art|num|aux|det|int|modal|part|phrasal|phr|n|v"
_PAREN_POS = re.compile(r"\(\s*(?:(?:%s)\.?\s*(?:[&,]\s*)?)+\)" % _POS, re.I)
_BARE_POS = re.compile(r"\s+((?:(?:%s)\.?\s*,?\s*)+)$" % _POS, re.I)

# Baglam kurulabilen turler (translate_words yalnizca bunlari kullanir).
_CONTEXT_POS = {"v", "n", "adj", "adv"}

# Tanim/baslik/footer gurultusu.
_NOISE = re.compile(
    r"©|page\s+\d+\s+of|/\s*\d+\s*$|langeek|cambridge english|oxford university|"
    r"vocabulary\s+list|wordlist|ucles|\.{4,}|^\s*no\.?\s+word|pronunciation|"
    r"definition|table of contents",
    re.I,
)
_IPA = re.compile(r"/[^/]{2,}/")  # telaffuz: /rɪˈleɪʃən/
_BULLETS = ("•", "●", "-", "*", "‣", "·")

# Tek basina bir tur kisaltmasi olan "kelimeler" (legend satirlarindan sizar).
_POS_WORDS = {p for p in _POS.split("|")} | {"mv", "aux", "modal"}


SUPPORTED_EXTENSIONS = (".txt", ".pdf", ".docx")


class ExtractionError(Exception):
    """Dosya okunamadi/ayristirilizamadi - kullaniciya gosterilebilir."""


def extract_entries(file_path):
    """Dosyadan (kelime, tur) ciftleri cikarir. tur: v|n|adj|adv veya None.

    .txt: her satir bir kelimedir (kullanici temiz liste hazirlamis) - tur
    isareti olmayan satirlar da gecerli kelimedir.
    .pdf/.docx: gurultulu kaynaklar. Tur isareti bir kelimenin varligini
    gosterir; isaretsiz satirlar (baslik, tanim, footer) atilir.
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".txt":
        return read_words_from_txt(file_path)
    elif ext == ".pdf":
        reader = _lines_from_pdf
    elif ext == ".docx":
        reader = _lines_from_docx
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    # Ucuncu parti kutuphaneler bozuk/gecersiz dosyalarda cesitli istisnalar
    # atar (PdfStreamError, PackageNotFoundError, ...). Hepsini tek bir
    # kullaniciya-gosterilebilir hataya sarmala.
    try:
        lines = reader(file_path)
    except Exception as e:
        raise ExtractionError(
            f"The {ext} file could not be read. Make sure it is a valid, "
            f"non-corrupted {ext} file."
        ) from e

    return [entry for line in lines if (entry := _word_from_line(line))]


# ── Satir kaynaklari ──────────────────────────────────────────────────────
def _lines_from_pdf(file_path):
    """pypdf'in metin cikarmasi tek-sutunlu listelerde (Oxford, B1) satir
    yapisini iyi koruyor. Cok-sutunlu tematik gridlerde (PET) sutunlar bir
    satira karisabilir; _word_from_line'daki token-sayisi ve gurultu filtreleri
    bu karisik satirlarin cogunu eler. Mukemmel degil ama en yaygin format olan
    duz listeler icin dogru."""
    from pypdf import PdfReader

    lines = []
    for page in PdfReader(file_path).pages:
        lines.extend((page.extract_text() or "").split("\n"))
    return lines


def _lines_from_docx(file_path):
    from docx import Document

    doc = Document(file_path)
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    # Tablo hucreleri (kelime listeleri sik sik tabloda olur).
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                lines.extend(l for l in cell.text.split("\n") if l.strip())
    return lines


# ── Satirdan kelime cikarma ───────────────────────────────────────────────
def _word_from_line(line):
    line = line.strip()
    if not line or _NOISE.search(line):
        return None
    if line[0] in _BULLETS:  # ornek cumle
        return None

    # Bastaki sira numarasi: "1 relation" / "23. spouse"
    line = re.sub(r"^\s*\d+[.)]?\s+", "", line)

    # Tur isaretini bul; kelime isaretten oncesi.
    pos = None
    m = _PAREN_POS.search(line)
    if m:
        word = line[: m.start()].strip()
        pos = _clean_pos(m.group(0))
    else:
        m = _BARE_POS.search(line)
        if not m:
            return None  # isaret yok -> guvenli tarafta at (baslik/gurultu olabilir)
        word = line[: m.start()].strip()
        pos = _clean_pos(m.group(1))

    word = _IPA.sub("", word).strip()
    # Aciklama parantezleri: "counter (long flat surface)" -> "counter",
    # "laboratory (lab)" -> "laboratory". POS parantezleri zaten yukarida ayrildi.
    word = re.sub(r"\s*\([^)]*\)", "", word).strip()

    # Dogrulama: makul bir kelime mi?
    if not word or not re.search(r"[A-Za-z]", word):
        return None
    if re.search(r"\d", word):  # numara sizmis
        return None
    if len(word.split()) > 4 or len(word) > 40:  # tanim/sutun karismasi sizmis
        return None
    if word.lower() in _POS_WORDS:  # legend satirindan gelen "conj"/"prep"
        return None

    return (word, pos)


def _clean_pos(marker):
    """'(v., n.)' -> 'v'  |  '(adj & prep)' -> 'adj'  |  '(phr.)' -> None"""
    first = re.split(r"[&,]", marker.strip().strip("()"))[0].strip().rstrip(".").lower()
    return first if first in _CONTEXT_POS else None
